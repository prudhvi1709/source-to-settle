#!/usr/bin/env python3
"""
Excel and Word Document Generator - Creates realistic business documents
"""

import random
from datetime import datetime, timedelta
from pathlib import Path
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, PieChart, Reference
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(42)

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"

print("="*80)
print("EXCEL & WORD DOCUMENT GENERATOR")
print("="*80)

# Load CSV data
print("\nLoading CSV data...")
vendors_df = pd.read_csv(DATA_DIR / "vendors.csv")
invoices_df = pd.read_csv(DATA_DIR / "invoices.csv")
po_df = pd.read_csv(DATA_DIR / "po_gr.csv")
supplier_df = pd.read_csv(DATA_DIR / "supplier_history.csv")

# ============================================================================
# EXCEL DOCUMENTS
# ============================================================================

print("\n[10/13] Generating Excel documents...")

# 1. Vendor Database Excel
def generate_vendor_database_excel():
    """Generate comprehensive vendor database workbook"""
    wb = Workbook()

    # Sheet 1: Vendor Master
    ws1 = wb.active
    ws1.title = "Vendor Master"

    # Headers
    headers = ['Vendor ID', 'Vendor Name', 'Country', 'Industry', 'Status', 'Risk Band', 'Email', 'Phone']
    ws1.append(headers)

    # Style headers
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    for col_num, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Add data
    for idx, row in vendors_df.iterrows():
        ws1.append([
            row['vendor_id'],
            row['vendor_name'],
            row['country'],
            row['industry'],
            row['status'],
            row['risk_band'],
            row['contact_email'],
            row['phone']
        ])

    # Auto-adjust column widths
    for column in ws1.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws1.column_dimensions[column_letter].width = min(max_length + 2, 50)

    # Sheet 2: Risk Summary
    ws2 = wb.create_sheet("Risk Summary")
    ws2.append(['Risk Band', 'Count', 'Percentage'])
    ws2.append([''] * 3)

    risk_summary = vendors_df['risk_band'].value_counts()
    for risk_band, count in risk_summary.items():
        percentage = (count / len(vendors_df)) * 100
        ws2.append([risk_band, count, f"{percentage:.1f}%"])

    # Apply styling
    for col in ws2['A1:C1']:
        for cell in col:
            cell.fill = header_fill
            cell.font = header_font

    # Sheet 3: Contact List
    ws3 = wb.create_sheet("Contact List")
    ws3.append(['Vendor Name', 'Primary Contact', 'Title', 'Email', 'Phone'])

    for idx, row in vendors_df.iterrows():
        ws3.append([
            row['vendor_name'],
            row['primary_contact_name'],
            row['primary_contact_title'],
            row['contact_email'],
            row['phone']
        ])

    # Style headers
    for col in ws3['A1:E1']:
        for cell in col:
            cell.fill = header_fill
            cell.font = header_font

    wb.save(DATA_DIR / "vendors_master_database.xlsx")
    print("   ✓ Generated vendor master database Excel")

generate_vendor_database_excel()

# 2. Supplier Scorecard Excel
def generate_supplier_scorecard(vendor_id):
    """Generate detailed supplier scorecard"""
    vendor_info = vendors_df[vendors_df['vendor_id'] == vendor_id].iloc[0]
    supplier_info = supplier_df[supplier_df['vendor_id'] == vendor_id]

    if len(supplier_info) == 0:
        return

    supplier_info = supplier_info.iloc[0]

    wb = Workbook()
    ws = wb.active
    ws.title = "Scorecard"

    # Title
    ws.merge_cells('A1:F1')
    ws['A1'] = f"SUPPLIER PERFORMANCE SCORECARD - {vendor_info['vendor_name']}"
    ws['A1'].font = Font(size=16, bold=True, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    # Basic info
    ws['A3'] = "Vendor ID:"
    ws['B3'] = vendor_id
    ws['A4'] = "Industry:"
    ws['B4'] = vendor_info['industry']
    ws['A5'] = "Risk Band:"
    ws['B5'] = vendor_info['risk_band']

    # Make labels bold
    for row in [3, 4, 5]:
        ws[f'A{row}'].font = Font(bold=True)

    # KPIs section
    ws['A7'] = "KEY PERFORMANCE INDICATORS"
    ws.merge_cells('A7:F7')
    ws['A7'].font = Font(size=14, bold=True, color="FFFFFF")
    ws['A7'].fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
    ws['A7'].alignment = Alignment(horizontal="center")

    # KPI Table
    kpi_headers = ['Metric', 'Value', 'Target', 'Status']
    ws.append([''] * 6)
    ws.append(kpi_headers + ['', ''])

    header_fill = PatternFill(start_color="A9D08E", end_color="A9D08E", fill_type="solid")
    for col in ws['A9:D9']:
        for cell in col:
            cell.fill = header_fill
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

    # KPI data
    kpis = [
        ('On-Time Payment Rate', f"{supplier_info['on_time_payment_rate']:.1f}%", "95%",
         "✓" if supplier_info['on_time_payment_rate'] >= 95 else "⚠"),
        ('Dispute Rate', f"{supplier_info['dispute_rate']:.1f}%", "<3%",
         "✓" if supplier_info['dispute_rate'] < 3 else "⚠"),
        ('Avg Cycle Time', f"{supplier_info['average_cycle_time_days']} days", "30 days",
         "✓" if supplier_info['average_cycle_time_days'] <= 30 else "⚠"),
        ('Quality Score', f"{supplier_info['quality_score']:.1f}", "85",
         "✓" if supplier_info['quality_score'] >= 85 else "⚠"),
        ('Delivery Score', f"{supplier_info['delivery_score']:.1f}", "90",
         "✓" if supplier_info['delivery_score'] >= 90 else "⚠"),
        ('Compliance Score', f"{supplier_info['compliance_score']:.1f}", "95",
         "✓" if supplier_info['compliance_score'] >= 95 else "⚠"),
    ]

    for kpi in kpis:
        ws.append(list(kpi) + ['', ''])

    # Financial Summary
    ws.append([''] * 6)
    ws.append(['FINANCIAL SUMMARY'] + [''] * 5)
    ws.merge_cells(f'A{ws.max_row}:F{ws.max_row}')
    ws[f'A{ws.max_row}'].font = Font(size=14, bold=True, color="FFFFFF")
    ws[f'A{ws.max_row}'].fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
    ws[f'A{ws.max_row}'].alignment = Alignment(horizontal="center")

    ws.append(['Total Invoices:', supplier_info['total_invoices_processed'], '', '', '', ''])
    ws.append(['Total Amount Paid:', f"₹{supplier_info['total_amount_paid']:,.2f}", '', '', '', ''])
    ws.append(['Last Invoice Date:', supplier_info['last_invoice_date'], '', '', '', ''])

    # Recommendation
    ws.append([''] * 6)
    ws.append(['RECOMMENDATION'] + [''] * 5)
    ws.merge_cells(f'A{ws.max_row}:F{ws.max_row}')
    ws[f'A{ws.max_row}'].font = Font(size=14, bold=True, color="FFFFFF")
    ws[f'A{ws.max_row}'].fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    ws[f'A{ws.max_row}'].alignment = Alignment(horizontal="center")

    ws.append([supplier_info['recommendation'], '', '', '', '', ''])
    ws[f'A{ws.max_row}'].font = Font(size=16, bold=True)

    # Adjust column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 10

    output_path = DATA_DIR / "supplier_performance" / f"{vendor_id}_scorecard.xlsx"
    wb.save(output_path)

# Generate scorecards for all vendors with history
for vendor_id in supplier_df['vendor_id'].tolist():
    generate_supplier_scorecard(vendor_id)

print(f"   ✓ Generated {len(supplier_df)} supplier scorecards")

# 3. Invoice Register Excel
def generate_invoice_register():
    """Generate invoice tracking register"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice Register"

    # Headers
    headers = ['Invoice ID', 'Vendor ID', 'Invoice Number', 'Date', 'Amount', 'Status', 'PO Ref', 'Payment Date']
    ws.append(headers)

    # Style headers
    header_fill = PatternFill(start_color="44546A", end_color="44546A", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Add invoice data
    for idx, row in invoices_df.iterrows():
        ws.append([
            row['invoice_id'],
            row['vendor_id'],
            row['invoice_number'],
            row['invoice_date'],
            row['total_amount'],
            row['status'],
            row['po_reference'] or 'N/A',
            row['payment_date'] if row['payment_date'] else 'Pending'
        ])

        # Color code by status
        current_row = ws.max_row
        status_cell = ws.cell(row=current_row, column=6)

        if row['status'] == 'MATCHED':
            status_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        elif row['status'] == 'EXCEPTION':
            status_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        elif row['status'] == 'DUPLICATE':
            status_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")

    # Auto-adjust columns
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = min(max_length + 2, 40)

    wb.save(DATA_DIR / "invoices_register.xlsx")
    print("   ✓ Generated invoice register Excel")

generate_invoice_register()

# ============================================================================
# WORD DOCUMENTS
# ============================================================================

print("\n[11/13] Generating Word documents...")

# 1. Contract Draft with Track Changes
def generate_contract_draft_docx(vendor_row):
    """Generate contract draft in Word format"""
    doc = Document()

    # Title
    title = doc.add_heading('MASTER SERVICE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"DRAFT - FOR REVIEW").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()

    # Contract metadata
    doc.add_paragraph(f"Contract Number: MSA-{vendor_row['vendor_id'].split('-')[1]}-{random.randint(1000,9999)}")
    doc.add_paragraph(f"Date: {vendor_row['onboarding_date']}")
    doc.add_paragraph(f"Version: 1.0 DRAFT")

    doc.add_paragraph()

    # Parties
    doc.add_heading('BETWEEN:', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Acme Corporation Pvt Ltd').bold = True
    p1.add_run(' ("Company")')
    doc.add_paragraph('123 Business Park, Mumbai - 400001')

    doc.add_heading('AND:', level=2)
    p2 = doc.add_paragraph()
    p2.add_run(vendor_row['vendor_name']).bold = True
    p2.add_run(' ("Vendor")')
    doc.add_paragraph(f"{vendor_row['city']}, {vendor_row['state'] or vendor_row['country']}")

    doc.add_page_break()

    # Terms
    doc.add_heading('1. SCOPE OF SERVICES', level=1)
    doc.add_paragraph(
        f"The Vendor agrees to provide {vendor_row['industry']} services to the Company "
        f"as detailed in the Statement of Work (SOW) attached hereto."
    )

    # Add comment simulation
    p = doc.add_paragraph()
    p.add_run('[LEGAL REVIEW: Please specify service deliverables more clearly]')
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    p.runs[0].italic = True

    doc.add_heading('2. PAYMENT TERMS', level=1)
    payment_terms = random.choice(['Net 30', 'Net 45', 'Net 60'])
    doc.add_paragraph(f"Payment terms: {payment_terms} days from invoice date.")
    doc.add_paragraph("Late payment will attract interest at 18% per annum.")

    doc.add_heading('3. TERM AND TERMINATION', level=1)
    doc.add_paragraph("This Agreement shall be valid for a period of 2 years from the effective date.")
    doc.add_paragraph("Either party may terminate this Agreement with 90 days written notice.")

    # Add revision note
    p = doc.add_paragraph()
    p.add_run('[PROCUREMENT: Consider reducing notice period to 60 days]')
    p.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    p.runs[0].italic = True

    doc.add_heading('4. LIABILITY AND INDEMNIFICATION', level=1)
    if vendor_row['risk_band'] == 'HIGH':
        p = doc.add_paragraph("The Vendor's liability shall be unlimited for all claims arising from this Agreement.")
        # Highlight risky clause
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.runs[0].bold = True

        p = doc.add_paragraph()
        p.add_run('[LEGAL: HIGH RISK - Unlimited liability clause needs revision]')
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.runs[0].italic = True
        p.runs[0].bold = True
    else:
        doc.add_paragraph("The Vendor's total liability shall be limited to the fees paid in the preceding 12 months.")

    doc.add_heading('5. CONFIDENTIALITY', level=1)
    doc.add_paragraph(
        "Both parties agree to maintain strict confidentiality of all proprietary and confidential "
        "information disclosed during the term of this Agreement."
    )

    doc.add_heading('6. GOVERNING LAW', level=1)
    doc.add_paragraph("This Agreement shall be governed by the laws of India.")
    doc.add_paragraph("Disputes shall be subject to the jurisdiction of courts in Mumbai.")

    doc.add_page_break()

    # Signatures
    doc.add_heading('SIGNATURES', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    table.cell(0, 0).text = 'For Company:'
    table.cell(0, 1).text = 'For Vendor:'
    table.cell(2, 0).text = '________________________'
    table.cell(2, 1).text = '________________________'
    table.cell(3, 0).text = 'Authorized Signatory'
    table.cell(3, 1).text = vendor_row['primary_contact_name']
    table.cell(4, 0).text = 'Date: ______________'
    table.cell(4, 1).text = f"Title: {vendor_row['primary_contact_title']}"

    # Save
    output_path = DATA_DIR / "contracts" / f"{vendor_row['vendor_id']}_contract_draft_v1.docx"
    doc.save(output_path)

# Generate contract drafts for approved vendors
approved_vendors = vendors_df[vendors_df['status'] == 'APPROVED']
for idx, vendor_row in approved_vendors.iterrows():
    generate_contract_draft_docx(vendor_row)

    if (idx + 1) % 5 == 0:
        print(f"   Generated {idx + 1}/{len(approved_vendors)} contract drafts...")

print(f"   ✓ Generated {len(approved_vendors)} contract draft Word documents")

# 2. Company Profile Document
def generate_company_profile(vendor_row):
    """Generate company profile document"""
    doc = Document()

    # Header
    header = doc.add_heading(vendor_row['vendor_name'], 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Company Profile')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    doc.add_paragraph()

    # Overview
    doc.add_heading('COMPANY OVERVIEW', level=1)
    doc.add_paragraph(
        f"{vendor_row['vendor_name']} is a leading {vendor_row['industry']} company "
        f"based in {vendor_row['city']}, {vendor_row['country']}. With a commitment to excellence "
        f"and innovation, we have been serving clients globally with best-in-class solutions."
    )

    # Key Information
    doc.add_heading('KEY INFORMATION', level=1)

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'

    info = [
        ('Registration Number', vendor_row['registration_number']),
        ('Industry', vendor_row['industry']),
        ('Country', vendor_row['country']),
        ('Website', vendor_row['website']),
        ('Email', vendor_row['contact_email']),
        ('Phone', vendor_row['phone']),
        ('PAN', vendor_row['pan'] or 'N/A'),
        ('GST', vendor_row['gst'] or 'N/A'),
    ]

    for idx, (label, value) in enumerate(info):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = str(value)
        table.cell(idx, 0).paragraphs[0].runs[0].bold = True

    # Services
    doc.add_heading('SERVICES OFFERED', level=1)
    services = [
        "End-to-end project delivery",
        "24/7 support and maintenance",
        "Custom solution development",
        "Consulting and advisory services",
        "Training and knowledge transfer"
    ]
    for service in services:
        doc.add_paragraph(service, style='List Bullet')

    # Certifications
    doc.add_heading('CERTIFICATIONS & COMPLIANCE', level=1)
    certs = ['ISO 9001:2015', 'ISO 27001:2013', 'CMMI Level 5']
    if vendor_row['risk_band'] == 'LOW':
        certs.extend(['SOC 2 Type II', 'GDPR Compliant'])

    for cert in certs:
        doc.add_paragraph(f"✓ {cert}", style='List Bullet')

    # Contact
    doc.add_heading('PRIMARY CONTACT', level=1)
    doc.add_paragraph(f"Name: {vendor_row['primary_contact_name']}")
    doc.add_paragraph(f"Title: {vendor_row['primary_contact_title']}")
    doc.add_paragraph(f"Email: {vendor_row['contact_email']}")

    output_path = DATA_DIR / "kyc_samples" / f"{vendor_row['vendor_id']}_company_profile.docx"
    doc.save(output_path)

# Generate company profiles
for idx, vendor_row in vendors_df.iterrows():
    generate_company_profile(vendor_row)

print(f"   ✓ Generated {len(vendors_df)} company profile Word documents")

print("\n" + "="*80)
print("EXCEL & WORD GENERATION COMPLETE!")
print("="*80)
